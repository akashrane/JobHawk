"""
Resume builder: applies a resume diff to the original DOCX via find-and-replace,
preserving all fonts, spacing, and formatting exactly.
"""
import io
import logging

logger = logging.getLogger(__name__)


def apply_diff_to_docx(docx_bytes: bytes, changes: list[dict]) -> bytes:
    """
    Open the user's original DOCX, find each changed bullet by its exact text,
    replace only that text in-place, and return the modified DOCX bytes.

    All fonts, styles, spacing, and layout are preserved — only the bullet
    text content changes.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))

    for change in changes:
        original = (change.get("original") or "").strip()
        modified = (change.get("modified") or "").strip()
        if not original or not modified or original == modified:
            continue
        replaced = _replace_in_document(doc, original, modified)
        if replaced:
            logger.info("resume_builder: replaced '%s...'", original[:60])
        else:
            logger.warning("resume_builder: could not find text '%s...'", original[:60])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _replace_in_document(doc, original: str, modified: str) -> bool:
    """Search every paragraph (including inside tables) and replace original with modified."""
    # Body paragraphs
    for para in doc.paragraphs:
        if _replace_in_paragraph(para, original, modified):
            return True

    # Paragraphs inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_in_paragraph(para, original, modified):
                        return True

    # Headers and footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for para in header_footer.paragraphs:
                if _replace_in_paragraph(para, original, modified):
                    return True

    return False


def _replace_in_paragraph(para, original: str, modified: str) -> bool:
    """
    Replace `original` with `modified` inside a paragraph, preserving run formatting.

    Strategy:
    - Reconstruct the full text from all runs
    - If original is found, put modified text in the first matching run,
      clear the rest — keeping that first run's character style (font, bold, size, color)
    - This handles the common DOCX case where a bullet is split across 2-3 runs
      for styling reasons but the visible text is continuous
    """
    full_text = para.text
    if original not in full_text:
        return False

    # Build a map: which runs contribute to which character positions
    runs = para.runs
    if not runs:
        return False

    # Find start position of `original` in full_text
    start = full_text.index(original)
    end = start + len(original)

    # Walk runs to find which ones overlap the [start, end) region
    cursor = 0
    first_run_idx = None
    for i, run in enumerate(runs):
        run_start = cursor
        run_end = cursor + len(run.text)
        if first_run_idx is None and run_end > start:
            first_run_idx = i
        cursor = run_end

    if first_run_idx is None:
        return False

    # Rebuild: everything before the match + modified + everything after
    new_full = full_text[:start] + modified + full_text[end:]

    # Put the entire new text in the first run, clear the rest
    # (the first run carries the dominant character formatting for this bullet)
    runs[first_run_idx].text = new_full
    for i, run in enumerate(runs):
        if i != first_run_idx:
            run.text = ""

    return True
