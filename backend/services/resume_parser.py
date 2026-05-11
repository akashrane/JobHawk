"""PDF extraction + LLM structuring for resumes."""
import io
import json
import logging

import pdfplumber
from pydantic import ValidationError

from db.models import ParsedResume
from services.llm import call_parsing_llm

logger = logging.getLogger(__name__)

PARSE_PROMPT = """Extract structured information from this resume text.
Return ONLY valid JSON matching this exact schema (no extra fields, no markdown):

{
  "name": "string",
  "email": "string or null",
  "phone": "string or null",
  "location": "string or null",
  "summary": "string or null",
  "experience": [
    {
      "company": "string",
      "title": "string",
      "start_date": "string (e.g. Jan 2022)",
      "end_date": "string or null (null means Present)",
      "bullets": ["string", ...]
    }
  ],
  "education": [
    {
      "institution": "string",
      "degree": "string",
      "field": "string or null",
      "graduation_date": "string or null",
      "gpa": "string or null"
    }
  ],
  "skills": ["string", ...],
  "projects": [
    {
      "name": "string",
      "description": "string or null",
      "technologies": ["string", ...],
      "url": "string or null"
    }
  ],
  "certifications": ["string", ...]
}

Resume text:
"""


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract plain text from PDF bytes using pdfplumber."""
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extract plain text from DOCX bytes.
    Captures body paragraphs, text boxes, headers, and footers so that
    names styled as large headings or placed in text boxes are not missed.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(docx_bytes))
    seen: set[str] = set()
    lines: list[str] = []

    def add(text: str) -> None:
        t = text.strip()
        if t and t not in seen:
            seen.add(t)
            lines.append(t)

    # 1. Headers and footers (name is often here in styled templates)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            for para in hf.paragraphs:
                add(para.text)

    # 2. Text boxes / drawing canvases (w:txbxContent)
    for txbx in doc.element.findall(".//" + qn("w:txbxContent")):
        for p_elem in txbx.findall(".//" + qn("w:p")):
            text = "".join(
                t.text for t in p_elem.findall(".//" + qn("w:t")) if t.text
            )
            add(text)

    # 3. Regular body paragraphs
    for para in doc.paragraphs:
        add(para.text)

    # 4. Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    add(para.text)

    return "\n".join(lines)


def parse_resume(file_bytes: bytes, mime_type: str = "application/pdf") -> tuple[ParsedResume, str]:
    """
    Parse a resume PDF or DOCX.
    Returns (ParsedResume, raw_text).
    Raises ValueError if parsing fails.
    """
    if mime_type == DOCX_MIME:
        raw_text = extract_text_from_docx(file_bytes)
    else:
        raw_text = extract_text_from_pdf(file_bytes)
    if not raw_text.strip():
        raise ValueError("Could not extract text from PDF. The file may be image-based or corrupted.")

    messages = [
        {"role": "system", "content": "You are a precise data extraction assistant. Extract resume data and return only valid JSON."},
        {"role": "user", "content": PARSE_PROMPT + raw_text[:8000]},  # Trim to avoid token limits
    ]

    parsed_json, _ = call_parsing_llm(messages)

    try:
        resume = ParsedResume(**parsed_json)
    except (ValidationError, TypeError) as exc:
        logger.error("Resume parsing validation failed: %s | raw json: %s", exc, parsed_json)
        raise ValueError(f"Resume parsing produced invalid structure: {exc}") from exc

    return resume, raw_text
